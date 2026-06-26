"""
Generates USER_MANUAL.docx from the manual content with embedded screenshots.
Run from the project root:  python scripts/build_manual_docx.py
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent.parent
IMG_DIR = ROOT / "frontend" / "public" / "mock"
OUT_FILE = ROOT / "USER_MANUAL.docx"

# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "CCCCCC")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.runs[0].font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)   # blue
    elif level == 2:
        p.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)   # dark
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x44, 0x55, 0x66)
    run.font.italic = True
    p.paragraph_format.space_after = Pt(4)
    return p


def add_screenshot(doc, filename, caption):
    img_path = IMG_DIR / filename
    if not img_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(5.8))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x66, 0x77)


def add_simple_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_border(table)
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        set_cell_bg(cell, "1D4ED8")
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "F8FAFC" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    # column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Inches(w)
    doc.add_paragraph()   # spacer
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, bold_part=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_part and text.startswith(bold_part):
        r1 = p.add_run(bold_part)
        r1.bold = True
        r2 = p.add_run(text[len(bold_part):])
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p


def page_break(doc):
    doc.add_page_break()


# ---------------------------------------------------------------------------
# build document
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    # -----------------------------------------------------------------------
    # TITLE PAGE
    # -----------------------------------------------------------------------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Al-Wakeel Al-Shamel")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Order Management System")
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("User Manual")
    r3.font.size = Pt(22)
    r3.bold = True
    r3.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc.add_paragraph()

    meta = [
        ("Version", "1.0"),
        ("Date", "June 2026"),
        ("Developer", "Osama Al-Hossam"),
        ("Website", "https://alwakeel-alshamel.vercel.app"),
    ]
    for k, v in meta:
        pm = doc.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rm = pm.add_run(f"{k}: ")
        rm.bold = True
        pm.add_run(v)

    page_break(doc)

    # -----------------------------------------------------------------------
    # 1. INTRODUCTION
    # -----------------------------------------------------------------------
    add_heading(doc, "1. Introduction", 1)

    add_heading(doc, "1.1 Purpose", 2)
    add_body(doc, (
        "This manual describes how to use the Al-Wakeel Al-Shamel Order Management System (OMS) "
        "— a web application for selling premium phone accessories online and managing orders, "
        "inventory, and business reports."
    ))

    add_heading(doc, "1.2 Intended Users", 2)
    add_simple_table(doc,
        ["Role", "Description"],
        [
            ["Customer", "Browses products, places orders, and manages their account"],
            ["Administrator", "Manages products, customers, orders, and sales reports"],
            ["Retail Salesperson", "Views sales performance and customer orders"],
            ["Warehouse Manager", "Manages stock levels and inventory check-ups"],
        ],
        col_widths=[1.8, 4.0],
    )

    add_heading(doc, "1.3 Live System Access", 2)
    add_simple_table(doc,
        ["Component", "URL"],
        [
            ["Website (Frontend)", "https://alwakeel-alshamel.vercel.app"],
            ["API (Backend)", "https://oms-api-23gt.onrender.com"],
        ],
        col_widths=[1.8, 4.0],
    )
    add_note(doc, (
        "Note: The backend runs on Render's free tier. If the site has been idle for 15+ minutes, "
        "the first action may take up to one minute while the server wakes up."
    ))

    page_break(doc)

    # -----------------------------------------------------------------------
    # 2. SYSTEM OVERVIEW
    # -----------------------------------------------------------------------
    add_heading(doc, "2. System Overview", 1)

    add_heading(doc, "2.1 Main Features", 2)
    features = [
        "Online product catalogue (chargers, earphones, power banks, phone cases)",
        "Shopping cart and secure checkout",
        "Payment by Credit Card or Cash on Delivery",
        "Automatic FedEx Express shipping on all orders",
        "Email verification for new customer accounts",
        "PDF invoices and order confirmation emails",
        "Role-based dashboards for Admin, Sales, and Warehouse staff",
        "Mobile-friendly layout with a collapsible navigation menu (☰)",
    ]
    for f in features:
        add_bullet(doc, f)

    add_heading(doc, "2.2 User Roles and Access", 2)
    add_body(doc, "Each user account is assigned one role. After login, you only see the menus and pages allowed for your role.")
    add_simple_table(doc,
        ["Role", "Pages Available"],
        [
            ["Customer", "Products, Cart, Orders, Profile"],
            ["Retail Salesperson", "Sales Dashboard, Orders"],
            ["Warehouse Manager", "Products (view), Inventory, Warehouse Dashboard"],
            ["Administrator", "Admin Panel: Dashboard, Products, Orders, Customers, Reports"],
        ],
        col_widths=[1.8, 4.0],
    )

    page_break(doc)

    # -----------------------------------------------------------------------
    # 3. GETTING STARTED
    # -----------------------------------------------------------------------
    add_heading(doc, "3. Getting Started", 1)

    add_heading(doc, "3.1 Supported Browsers and Devices", 2)
    add_bullet(doc, "Browsers: Chrome, Edge, Firefox, Safari (latest versions recommended)")
    add_bullet(doc, "Devices: Desktop, laptop, tablet, and smartphone")
    add_bullet(doc, "On mobile, tap the ☰ menu icon in the top bar to open navigation links")

    add_heading(doc, "3.2 Demo Accounts (Testing)", 2)
    add_body(doc, "The system includes pre-configured demo accounts. All use the password: DemoPass!123")
    add_simple_table(doc,
        ["Role", "Email", "Login Page"],
        [
            ["Customer", "customer@demo.local", "Main Login"],
            ["Administrator", "admin@demo.local", "Admin Login (/admin/login)"],
            ["Retail Salesperson", "sales@demo.local", "Main Login"],
            ["Warehouse Manager", "warehouse@demo.local", "Main Login"],
        ],
        col_widths=[1.6, 2.2, 2.0],
    )

    add_heading(doc, "3.3 Logging In", 2)
    steps = [
        "Open the website: https://alwakeel-alshamel.vercel.app",
        "Click Login in the top-right corner.",
        "Enter your email and password.",
        "Click Sign In.",
        "You are redirected to your role's home page (e.g. Products for customers).",
    ]
    for s in steps:
        add_numbered(doc, s)
    doc.add_paragraph()
    add_screenshot(doc, "frame-14.png", "Figure 1 — Login Page")
    add_note(doc, "Administrators must use the separate admin login at /admin/login. Click 'Admin Login →' at the bottom of the main login page.")

    add_heading(doc, "3.4 Logging Out", 2)
    add_body(doc, "Click Logout in the top-right corner. You will be returned to the login page.")

    add_heading(doc, "3.5 Creating a New Customer Account", 2)
    reg_steps = [
        "Click Sign Up on the home page or login page.",
        "Fill in your full name, phone number, email address, delivery address, and password.",
        "Click Create Account.",
        "A verification email is sent to your inbox.",
        "Open the email and click the verification link.",
        "After verification, log in and start shopping.",
    ]
    for s in reg_steps:
        add_numbered(doc, s)
    doc.add_paragraph()
    add_screenshot(doc, "frame-05.png", "Figure 2 — Create Account form")
    add_note(doc, "Important: You must verify your email before you can use the cart or place orders. If the email does not arrive, check your spam folder or resend from Profile.")

    page_break(doc)

    # -----------------------------------------------------------------------
    # 4. CUSTOMER GUIDE
    # -----------------------------------------------------------------------
    add_heading(doc, "4. Customer Guide", 1)

    add_heading(doc, "4.1 Home Page", 2)
    add_body(doc, "The home page introduces Al-Wakeel Al-Shamel and provides:")
    add_bullet(doc, "A Browse Products button (redirects to login if not signed in)")
    add_bullet(doc, "A Sign Up button for new visitors")
    add_bullet(doc, "A 'Why Shop With Us?' section highlighting key benefits")
    add_bullet(doc, "Popular Categories — visual links to Chargers, Earphones, Power Banks, and Phone Cases")
    doc.add_paragraph()
    add_screenshot(doc, "frame-03.png", "Figure 3 — Home Page hero banner")
    add_screenshot(doc, "frame-13.png", "Figure 4 — Popular Categories section")
    add_screenshot(doc, "frame-08.png", "Figure 5 — Call-to-action and footer")

    add_heading(doc, "4.2 Browsing Products", 2)
    browse_steps = [
        "After logging in, click Products in the navigation bar.",
        "Use the search bar to find products by name.",
        "Use the category filter dropdown (All, Chargers, Earphones, etc.) to narrow results.",
        "Each product card shows the image, name, description, price, and stock level.",
        "Click any product card to open the Product Details page.",
        "On the details page, review the description, price, stock level, and product features.",
    ]
    for s in browse_steps:
        add_numbered(doc, s)
    doc.add_paragraph()
    add_screenshot(doc, "frame-15.png", "Figure 6 — Products Page with search and category filter")
    add_screenshot(doc, "frame-01.png", "Figure 7 — Product Details page")

    add_heading(doc, "4.3 Adding Items to the Cart", 2)
    add_bullet(doc, "On the Products page, click Add to cart on any in-stock product card.")
    add_bullet(doc, "Or open a product's detail page, set the quantity, and click Add to Cart.")
    add_bullet(doc, "The cart icon (🛒) in the top bar shows a badge with the total number of items.")

    add_heading(doc, "4.4 Managing the Shopping Cart", 2)
    add_body(doc, "Click the cart icon in the top bar to open the Shopping Cart page.")
    add_screenshot(doc, "frame-07.png", "Figure 8 — Shopping Cart with Order Summary panel")
    add_simple_table(doc,
        ["Action", "How"],
        [
            ["Change quantity", "Use the − and + buttons next to an item"],
            ["Remove an item", "Click the 🗑 (trash) icon"],
            ["Review totals", "See subtotal, FedEx shipping (Free), and total in the Order Summary panel"],
        ],
        col_widths=[1.8, 4.0],
    )
    add_note(doc, "Shipping: All orders ship via FedEx Express (estimated 3–5 business days). Shipping is free on all orders.")

    add_heading(doc, "4.5 Checkout and Payment", 2)
    add_body(doc, "From the Cart page, select a payment method and complete checkout. The process follows four steps shown at the top of the page:")
    add_code_block(doc, "Cart  →  Payment Verification  →  Order Confirmation  →  Invoice")

    p = doc.add_paragraph()
    p.add_run("Step 1 — Cart (Order Details)").bold = True
    checkout_steps = [
        "Select a Payment Method: Credit Card (pay online now) or Cash on Delivery (pay on arrival).",
        "If paying by Credit Card, enter: Cardholder name, Card number, Expiry date (MM/YY), and CVV.",
        "Click Proceed to Payment Verification.",
    ]
    for s in checkout_steps:
        add_numbered(doc, s)
    add_note(doc, "For testing, use card number 4111 1111 1111 1111 with any valid future expiry date and a 3-digit CVV.")

    p = doc.add_paragraph()
    p.add_run("Step 2 — Payment Verification").bold = True
    add_body(doc, "For card payments, a short verification screen simulates payment processing. Wait for the countdown or click Proceed when ready. For Cash on Delivery, this step confirms your order details before finalising.")

    p = doc.add_paragraph()
    p.add_run("Step 3 — Order Confirmation").bold = True
    add_body(doc, "You see your Order ID, total amount, payment method, and FedEx tracking number. A confirmation email with a PDF invoice is sent to your registered email address.")

    p = doc.add_paragraph()
    p.add_run("Step 4 — Invoice").bold = True
    add_body(doc, "View the full invoice on screen. Options: Download PDF, Print, or return to Order History.")

    add_heading(doc, "4.6 Order History", 2)
    order_steps = [
        "Click Orders in the navigation bar.",
        "View all past orders with order ID, date, status, items, prices, and FedEx tracking number.",
        "Filter by status using the dropdown at the top (Placed, Processing, Shipped, Delivered, Cancelled).",
    ]
    for s in order_steps:
        add_numbered(doc, s)
    doc.add_paragraph()
    add_screenshot(doc, "frame-06.png", "Figure 9 — Order History with status badges and item breakdown")
    add_body(doc, "For each order you can:")
    add_bullet(doc, "View invoice — opens the invoice page")
    add_bullet(doc, "Download PDF — saves the invoice as a PDF file")
    add_bullet(doc, "Reorder — copies items back into your cart")

    add_heading(doc, "4.7 My Profile", 2)
    profile_steps = [
        "Click Profile in the navigation bar.",
        "View your account details: name, email, phone, address, and member-since date.",
        "The Account Statistics section shows total orders, total amount spent, completed orders, and pending orders.",
        "Edit profile: Update your name, phone, or address → click Save Changes.",
        "Change password: Enter your current and new password → click Update Password.",
        "Email verification: If your email is not yet verified, click Resend verification email.",
    ]
    for s in profile_steps:
        add_numbered(doc, s)
    doc.add_paragraph()
    add_screenshot(doc, "frame-04.png", "Figure 10 — My Profile page with Account Statistics")

    page_break(doc)

    # -----------------------------------------------------------------------
    # 5. ADMINISTRATOR GUIDE
    # -----------------------------------------------------------------------
    add_heading(doc, "5. Administrator Guide", 1)
    add_body(doc, "Administrators use a separate dark-themed Admin Panel with its own navigation bar.")

    add_heading(doc, "5.1 Accessing the Admin Panel", 2)
    admin_steps = [
        "Go to https://alwakeel-alshamel.vercel.app/admin/login",
        "Log in with an Admin account (e.g. admin@demo.local / DemoPass!123).",
        "You are taken to the Admin Dashboard.",
    ]
    for s in admin_steps:
        add_numbered(doc, s)

    add_heading(doc, "5.2 Admin Dashboard", 2)
    add_body(doc, "The dashboard provides a real-time overview of the entire business.")
    add_screenshot(doc, "frame-12.png", "Figure 11 — Admin Dashboard with KPI cards, Recent Orders, and Low Stock Alert")
    add_simple_table(doc,
        ["Card", "Description"],
        [
            ["Total Orders", "All orders in the system; pending count shown in amber"],
            ["Total Products", "Number of active products in the catalogue"],
            ["Total Customers", "Number of registered customer accounts"],
            ["Total Revenue", "Cumulative cash revenue across all orders"],
            ["Recent Orders", "Latest orders with customer name, status badge, and total"],
            ["Low Stock Alert", "Products at or near minimum stock threshold"],
        ],
        col_widths=[1.8, 4.0],
    )
    add_note(doc, "Click Refresh to reload the latest data from the server.")

    add_heading(doc, "5.3 Product Management", 2)
    add_body(doc, "Navigation: Admin Panel → Products")
    add_screenshot(doc, "frame-02.png", "Figure 12 — Admin Product Management with edit and delete actions")
    add_simple_table(doc,
        ["Task", "Steps"],
        [
            ["View all products", "Scroll the product table; use the search box to filter by name or ID"],
            ["Add a product", "Click + Add Product, fill in the form (ID, Name, Category, Price, Stock, Description, Image URL), then click Save product"],
            ["Edit a product", "Click the ✎ (pencil) icon on a row — the form fills with current data. Modify fields, then click Save product"],
            ["Delete a product", "Click the 🗑 icon and confirm the deletion"],
        ],
        col_widths=[1.6, 4.2],
    )
    add_note(doc, "Image URL tip: Use paths like /images/products/charger.jpg for images stored in the public folder, or supply a full external URL.")

    add_heading(doc, "5.4 Order Management", 2)
    add_body(doc, "Navigation: Admin Panel → Orders")
    add_screenshot(doc, "frame-09.png", "Figure 13 — Admin Order Management table")
    add_bullet(doc, "View all customer orders in a searchable table.")
    add_bullet(doc, "Columns: Order ID, Customer (name and ID), Date, Total, Payment Method, Shipping (FedEx tracking and ETA).")
    add_bullet(doc, "Use the search box to find orders by Order ID or customer email.")
    add_note(doc, "Orders are fulfilled automatically via FedEx after checkout — no manual status changes are required.")

    add_heading(doc, "5.5 Customer Management", 2)
    add_body(doc, "Navigation: Admin Panel → Customers")
    add_screenshot(doc, "frame-16.png", "Figure 14 — Admin Customer Management table")
    add_bullet(doc, "View all registered customers with name, customer ID, email, phone number, registration date, and order count.")
    add_bullet(doc, "Use the search box to find a specific customer by name or email.")

    add_heading(doc, "5.6 Sales Reports", 2)
    add_body(doc, "Navigation: Admin Panel → Reports")
    add_screenshot(doc, "frame-11.png", "Figure 15 — Sales Report KPI cards")
    add_simple_table(doc,
        ["Metric", "Description"],
        [
            ["Total Revenue", "Gross revenue across all orders"],
            ["Completed Revenue", "Revenue from fully completed orders only"],
            ["Total Orders", "Total number of orders placed"],
            ["Avg Order Value", "Average spend per order"],
        ],
        col_widths=[1.8, 4.0],
    )
    add_body(doc, "Scrolling down shows detailed breakdowns:")
    add_screenshot(doc, "frame-10.png", "Figure 16 — Daily Sales Summary and Order Status Breakdown")
    add_bullet(doc, "Daily Sales Summary — orders and revenue per day, sorted by most recent date.")
    add_bullet(doc, "Order Status Breakdown — count and revenue split by Pending, Confirmed, Completed, and Cancelled.")
    add_note(doc, "Click Export Report to download the report data for offline use.")

    page_break(doc)

    # -----------------------------------------------------------------------
    # 6. RETAIL SALESPERSON GUIDE
    # -----------------------------------------------------------------------
    add_heading(doc, "6. Retail Salesperson Guide", 1)

    add_heading(doc, "6.1 Sales Dashboard", 2)
    add_body(doc, "Navigation: After login → automatically redirected to the Sales Dashboard.")
    add_body(doc, "The Sales Dashboard shows:")
    add_bullet(doc, "Total orders, cash orders, and cash revenue metrics")
    add_bullet(doc, "A table of recent customer orders (Order ID, customer email, status, total, and date)")

    add_heading(doc, "6.2 Viewing Orders", 2)
    add_body(doc, "Navigation: Orders in the top menu.")
    add_bullet(doc, "Read-only list of all customer orders.")
    add_bullet(doc, "Displays payment method, FedEx shipping details, and order totals.")
    add_bullet(doc, "Use this page to answer customer enquiries and monitor sales activity.")
    add_note(doc, "Sales staff cannot modify orders or approve payments — all orders are processed automatically at checkout.")

    page_break(doc)

    # -----------------------------------------------------------------------
    # 7. WAREHOUSE MANAGER GUIDE
    # -----------------------------------------------------------------------
    add_heading(doc, "7. Warehouse Manager Guide", 1)

    add_heading(doc, "7.1 Warehouse Dashboard", 2)
    add_body(doc, "Navigation: After login → Dashboard.")
    add_body(doc, "The Warehouse Dashboard shows:")
    add_bullet(doc, "Low Stock alerts — products at or below the stock threshold, sorted by quantity ascending")
    add_bullet(doc, "Recent Inventory Activity — a timestamped log of stock received, check-ups, and order deductions")

    add_heading(doc, "7.2 Inventory Management", 2)
    add_body(doc, "Navigation: Inventory in the top menu. The page has two action panels at the top and a full inventory table below.")

    p = doc.add_paragraph()
    p.add_run("Receive Stock").bold = True
    receive_steps = [
        "Select a product from the dropdown list.",
        "Enter the quantity received (minimum 1).",
        "Add an optional note (e.g. 'Supplier delivery #4521').",
        "Click Receive.",
        "The product's available quantity and stock level increase automatically.",
    ]
    for s in receive_steps:
        add_numbered(doc, s)

    p = doc.add_paragraph()
    p.add_run("Log Check-up").bold = True
    checkup_steps = [
        "Select an inventory location from the dropdown list.",
        "Add an optional note describing the check-up findings.",
        "Click Log check-up.",
        "The Last Check-up date for that inventory row is updated.",
    ]
    for s in checkup_steps:
        add_numbered(doc, s)

    add_heading(doc, "Inventory Table Columns", 3)
    add_simple_table(doc,
        ["Column", "Description"],
        [
            ["Product", "Product name and Product ID"],
            ["Location", "Warehouse location code"],
            ["Available", "Units currently available for sale"],
            ["Reserved", "Units reserved for pending orders"],
            ["Last Check-up", "Date and time of the most recent check-up"],
        ],
        col_widths=[1.6, 4.2],
    )

    add_heading(doc, "7.3 Viewing Products", 2)
    add_body(doc, "Warehouse managers can open the Products page to view the full catalogue and current stock levels (read-only). Product editing is performed by Administrators only.")

    page_break(doc)

    # -----------------------------------------------------------------------
    # 8. TROUBLESHOOTING
    # -----------------------------------------------------------------------
    add_heading(doc, "8. Troubleshooting", 1)

    add_heading(doc, "8.1 Common Issues", 2)
    add_simple_table(doc,
        ["Problem", "Likely Cause", "Solution"],
        [
            ["Page loads slowly on first visit", "API server waking from sleep (Render free tier)", "Wait 30–60 seconds and refresh"],
            ["'Failed to load profile' or blank data", "Stale login token", "Log out and log in again"],
            ["Cannot access Cart or Checkout", "Email not verified", "Click the link in your verification email, or resend it from Profile"],
            ["Login fails for demo accounts", "Typo in email or password", "Use exact emails from Section 3.2 with password DemoPass!123"],
            ["Credit card payment rejected", "Invalid test card details", "Use 4111 1111 1111 1111 with a future expiry date and any 3-digit CVV"],
            ["No verification email received", "Email may be in spam folder", "Check spam/junk; contact the administrator if still missing"],
            ["Products show zero stock", "Inventory depleted", "Warehouse Manager receives new stock, or Admin increases the stock level"],
            ["Navigation not visible on mobile", "Menu is collapsed", "Tap the ☰ icon in the top bar to expand the menu"],
        ],
        col_widths=[1.8, 1.8, 2.2],
    )

    add_heading(doc, "8.2 Clearing Browser Session", 2)
    add_body(doc, "If you experience persistent errors after a system update:")
    for s in ["Click Logout.", "Clear your browser cache, or open the site in a private/incognito window.", "Log in again."]:
        add_numbered(doc, s)

    add_heading(doc, "8.3 Getting Help", 2)
    add_body(doc, "For technical support, contact the system developer or your course instructor and include:")
    for item in [
        "Your role and the email address used to log in",
        "The page where the error occurred",
        "A screenshot of any error message displayed",
        "The approximate date and time of the issue",
    ]:
        add_bullet(doc, item)

    page_break(doc)

    # -----------------------------------------------------------------------
    # 9. APPENDIX
    # -----------------------------------------------------------------------
    add_heading(doc, "9. Appendix", 1)

    add_heading(doc, "9.1 Glossary", 2)
    add_simple_table(doc,
        ["Term", "Definition"],
        [
            ["OMS", "Order Management System"],
            ["SKU / Product ID", "Unique code for each product (e.g. P001)"],
            ["COD", "Cash on Delivery — payment collected when the order arrives"],
            ["FedEx Express", "Courier service used for all shipments in this system"],
            ["JWT", "Secure token used to keep the user logged in between pages"],
            ["Invoice", "PDF document listing order items, totals, and payment details"],
        ],
        col_widths=[1.6, 4.2],
    )

    add_heading(doc, "9.2 Pre-loaded Demo Products", 2)
    add_simple_table(doc,
        ["Product ID", "Name", "Category"],
        [
            ["P001", "Fast Charging Cable USB-C", "Chargers"],
            ["P002", "Wireless Charging Pad", "Chargers"],
            ["P003", "True Wireless Earbuds", "Earphones"],
            ["P004", "Sport Earphones", "Earphones"],
        ],
        col_widths=[1.0, 2.6, 1.6],
    )

    add_heading(doc, "9.3 Checkout Flow Diagram", 2)
    add_code_block(doc,
        "┌─────────┐   ┌──────────────────────┐   ┌─────────────────┐   ┌─────────┐\n"
        "│  Cart   │ → │ Payment Verification │ → │  Confirmation   │ → │ Invoice │\n"
        "│         │   │  (card or COD)       │   │  + Email sent   │   │  (PDF)  │\n"
        "└─────────┘   └──────────────────────┘   └─────────────────┘   └─────────┘\n"
        "                       ↓\n"
        "              FedEx order created\n"
        "              Tracking number assigned"
    )

    add_heading(doc, "9.4 Document Revision History", 2)
    add_simple_table(doc,
        ["Version", "Date", "Changes"],
        [["1.0", "June 2026", "Initial user manual for production deployment"]],
        col_widths=[0.8, 1.2, 3.8],
    )

    # final note
    doc.add_paragraph()
    p = doc.add_paragraph("End of User Manual")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x88, 0x99, 0xAA)

    # -----------------------------------------------------------------------
    doc.save(str(OUT_FILE))
    print(f"Saved: {OUT_FILE}")


if __name__ == "__main__":
    build()
